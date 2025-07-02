from docx import Document
import os

def write_to_word(module_title, steps, status, notes=""):
    filepath = r"D:\Desktop\selenium_testing\selenium_multibrowser_project\docs\Project_Report.docx"
    folder = os.path.dirname(filepath)
    if not os.path.exists(folder):
        os.makedirs(folder)
    if os.path.exists(filepath):
        doc = Document(filepath)
    else:
        doc = Document()
        doc.add_heading("Selenium Multi-Browser Automation Report", level=1)
        doc.save(filepath)  # Save immediately to create file
    doc.add_heading(module_title, level=2)
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    doc.add_paragraph(f"Status: {status}", style='Intense Quote')
    if notes:
        doc.add_paragraph(f"Notes: {notes}", style='Normal')
    doc.add_paragraph("\n")
    doc.save(filepath)
