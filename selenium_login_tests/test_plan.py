from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from datetime import datetime

def create_login_test_plan():
    """
    Creates a comprehensive test plan document for login functionality
    Requires: pip install python-docx
    """
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Login Functionality Test Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document information
    doc.add_heading('Document Information', level=1)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ['Document Title', 'Login Functionality Test Plan'],
        ['Application', 'SauceDemo Web Application'],
        ['Version', '1.0'],
        ['Created By', 'QA Team'],
        ['Creation Date', datetime.now().strftime('%Y-%m-%d')],
        ['Status', 'Draft']
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
    
    # Test Objectives
    doc.add_heading('Test Objectives', level=1)
    objectives = [
        'Verify login functionality with valid credentials',
        'Validate error handling for invalid credentials',
        'Test field validation and user experience',
        'Ensure security measures are working correctly',
        'Confirm proper navigation after successful login'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # Test Scope
    doc.add_heading('Test Scope', level=1)
    
    doc.add_heading('In Scope:', level=2)
    in_scope = [
        'Username and password field validation',
        'Login button functionality',
        'Error message display and content',
        'Successful login navigation',
        'Form field clearing and input handling'
    ]
    
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Out of Scope:', level=2)
    out_scope = [
        'Password reset functionality',
        'User registration process',
        'Session management and timeout',
        'Database connectivity testing',
        'Performance and load testing'
    ]
    
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    # Test Environment
    doc.add_heading('Test Environment', level=1)
    env_table = doc.add_table(rows=5, cols=2)
    env_table.style = 'Table Grid'
    
    env_data = [
        ['Application URL', 'https://www.saucedemo.com/'],
        ['Browser', 'Chrome (Latest Version)'],
        ['Operating System', 'Windows 10/11'],
        ['Test Framework', 'Selenium WebDriver with Python'],
        ['Test Runner', 'Pytest']
    ]
    
    for i, (key, value) in enumerate(env_data):
        env_table.cell(i, 0).text = key
        env_table.cell(i, 1).text = value
    
    # Test Data
    doc.add_heading('Test Data', level=1)
    
    doc.add_heading('Valid Credentials:', level=2)
    valid_table = doc.add_table(rows=2, cols=2)
    valid_table.style = 'Table Grid'
    valid_table.cell(0, 0).text = 'Username'
    valid_table.cell(0, 1).text = 'Password'
    valid_table.cell(1, 0).text = 'standard_user'
    valid_table.cell(1, 1).text = 'secret_sauce'
    
    doc.add_heading('Invalid Test Data:', level=2)
    invalid_data = [
        'Invalid username: "invalid_user"',
        'Invalid password: "wrong_password"',
        'Empty username: ""',
        'Empty password: ""',
        'Special characters and SQL injection attempts'
    ]
    
    for data in invalid_data:
        doc.add_paragraph(data, style='List Bullet')
    
    # Test Cases
    doc.add_heading('Test Cases', level=1)
    
    # Create test cases table
    test_cases = [
        {
            'id': 'TC001',
            'title': 'Valid Credentials Login',
            'description': 'Verify successful login with valid username and password',
            'preconditions': 'User is on login page',
            'steps': '1. Enter valid username\n2. Enter valid password\n3. Click login button',
            'expected': 'User successfully logged in and redirected to dashboard',
            'priority': 'High'
        },
        {
            'id': 'TC002',
            'title': 'Invalid Username Login',
            'description': 'Verify login fails with invalid username',
            'preconditions': 'User is on login page',
            'steps': '1. Enter invalid username\n2. Enter valid password\n3. Click login button',
            'expected': 'Error message displayed: "Username and password do not match"',
            'priority': 'High'
        },
        {
            'id': 'TC003',
            'title': 'Invalid Password Login',
            'description': 'Verify login fails with invalid password',
            'preconditions': 'User is on login page',
            'steps': '1. Enter valid username\n2. Enter invalid password\n3. Click login button',
            'expected': 'Error message displayed: "Username and password do not match"',
            'priority': 'High'
        },
        {
            'id': 'TC004',
            'title': 'Empty Username Field',
            'description': 'Verify validation for empty username field',
            'preconditions': 'User is on login page',
            'steps': '1. Leave username field empty\n2. Enter valid password\n3. Click login button',
            'expected': 'Error message displayed: "Username is required"',
            'priority': 'Medium'
        },
        {
            'id': 'TC005',
            'title': 'Empty Password Field',
            'description': 'Verify validation for empty password field',
            'preconditions': 'User is on login page',
            'steps': '1. Enter valid username\n2. Leave password field empty\n3. Click login button',
            'expected': 'Error message displayed: "Password is required"',
            'priority': 'Medium'
        },
        {
            'id': 'TC006',
            'title': 'Both Fields Empty',
            'description': 'Verify validation when both fields are empty',
            'preconditions': 'User is on login page',
            'steps': '1. Leave both fields empty\n2. Click login button',
            'expected': 'Error message displayed: "Username is required"',
            'priority': 'Medium'
        },
        {
            'id': 'TC007',
            'title': 'Field Clearing Functionality',
            'description': 'Verify fields can be cleared and re-entered',
            'preconditions': 'User is on login page',
            'steps': '1. Enter text in both fields\n2. Clear fields\n3. Enter new text',
            'expected': 'Fields should clear completely and accept new input',
            'priority': 'Low'
        }
    ]
    
    # Create detailed test case table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    headers = ['Test ID', 'Test Title', 'Description', 'Preconditions', 'Test Steps', 'Expected Result', 'Priority']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        # Make header bold
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # Add test cases
    for test_case in test_cases:
        row = table.add_row()
        row.cells[0].text = test_case['id']
        row.cells[1].text = test_case['title']
        row.cells[2].text = test_case['description']
        row.cells[3].text = test_case['preconditions']
        row.cells[4].text = test_case['steps']
        row.cells[5].text = test_case['expected']
        row.cells[6].text = test_case['priority']
    
    # Set column widths
    for i, width in enumerate([0.8, 1.5, 2.0, 1.2, 2.0, 2.0, 0.8]):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # Test Execution Strategy
    doc.add_heading('Test Execution Strategy', level=1)
    
    doc.add_heading('Automation Approach:', level=2)
    automation_points = [
        'Use Selenium WebDriver for browser automation',
        'Implement Page Object Model for maintainability',
        'Use pytest framework for test execution and reporting',
        'Implement proper waits and error handling',
        'Generate detailed test reports with screenshots'
    ]
    
    for point in automation_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('Test Execution Order:', level=2)
    execution_order = [
        'Execute positive test cases first (TC001)',
        'Run negative test cases (TC002, TC003)',
        'Validate field validation tests (TC004, TC005, TC006)',
        'Execute UI functionality tests (TC007)',
        'Perform regression testing after any changes'
    ]
    
    for order in execution_order:
        doc.add_paragraph(order, style='List Bullet')
    
    # Risk Assessment
    doc.add_heading('Risk Assessment', level=1)
    risk_table = doc.add_table(rows=4, cols=3)
    risk_table.style = 'Table Grid'
    
    risk_headers = ['Risk', 'Impact', 'Mitigation']
    for i, header in enumerate(risk_headers):
        risk_table.cell(0, i).text = header
        risk_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    risks = [
        ['Application unavailable during testing', 'High', 'Have backup test environment'],
        ['Test data becomes invalid', 'Medium', 'Maintain updated test data repository'],
        ['Browser compatibility issues', 'Low', 'Test on multiple browsers if needed']
    ]
    
    for i, (risk, impact, mitigation) in enumerate(risks, 1):
        risk_table.cell(i, 0).text = risk
        risk_table.cell(i, 1).text = impact
        risk_table.cell(i, 2).text = mitigation
    
    # Deliverables
    doc.add_heading('Deliverables', level=1)
    deliverables = [
        'Test plan document (this document)',
        'Automated test scripts',
        'Test execution reports',
        'Defect reports (if any)',
        'Test summary report'
    ]
    
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    # Timeline
    doc.add_heading('Timeline', level=1)
    timeline_table = doc.add_table(rows=5, cols=3)
    timeline_table.style = 'Table Grid'
    
    timeline_headers = ['Phase', 'Duration', 'Deliverable']
    for i, header in enumerate(timeline_headers):
        timeline_table.cell(0, i).text = header
        timeline_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    timeline_data = [
        ['Test Planning', '1 day', 'Test Plan Document'],
        ['Test Script Development', '2 days', 'Automated Test Scripts'],
        ['Test Execution', '1 day', 'Test Results'],
        ['Reporting', '0.5 days', 'Final Test Report']
    ]
    
    for i, (phase, duration, deliverable) in enumerate(timeline_data, 1):
        timeline_table.cell(i, 0).text = phase
        timeline_table.cell(i, 1).text = duration
        timeline_table.cell(i, 2).text = deliverable
    
    # Add footer
    doc.add_page_break()
    doc.add_heading('Appendix', level=1)
    doc.add_paragraph('This test plan should be reviewed and approved by the project stakeholders before test execution begins.')
    
    # Save the document
    filename = f'Login_Test_Plan_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    
    return filename

# Usage example
if __name__ == "__main__":
    try:
        filename = create_login_test_plan()
        
    except ImportError:
        print("Error: python-docx package not installed")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating test plan: {str(e)}")

def create_custom_test_plan(app_name, test_cases_data):
    """
    Create a custom test plan with your own data
    
    Args:
        app_name (str): Name of the application being tested
        test_cases_data (list): List of dictionaries containing test case information
    
    Returns:
        str: Filename of the created document
    """
    doc = Document()
    
    # Add custom title
    title = doc.add_heading(f'{app_name} Test Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add test cases from custom data
    doc.add_heading('Test Cases', level=1)
    
    if test_cases_data:
        table = doc.add_table(rows=1, cols=len(test_cases_data[0].keys()))
        table.style = 'Table Grid'
        
        # Add headers
        headers = list(test_cases_data[0].keys())
        for i, header in enumerate(headers):
            table.cell(0, i).text = header.title()
            table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Add test case data
        for test_case in test_cases_data:
            row = table.add_row()
            for i, (key, value) in enumerate(test_case.items()):
                row.cells[i].text = str(value)
    
    # Save custom document
    filename = f'{app_name}_Test_Plan_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    
    return filename